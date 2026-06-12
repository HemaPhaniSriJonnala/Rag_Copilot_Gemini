"""
DocumentStore — SQLite-backed storage for documents and chunks.

Persists across server restarts (Render free tier spin-downs, etc.).
SQLite is part of Python stdlib — no extra install needed.

API is identical to the old in-memory version so main.py needs no changes.
"""

import uuid
import sqlite3
import os
import io
from dataclasses import dataclass
from datetime import datetime

DB_PATH = os.getenv("DOCS_DB_PATH", "documents.db")


# ══════════════════════════════════════════════════════════════════════════════
# Data models
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class Chunk:
    id:          str
    doc_id:      str
    doc_name:    str
    chunk_index: int
    text:        str

    @property
    def index(self) -> int:
        return self.chunk_index


@dataclass
class Document:
    id:          str
    name:        str
    preview:     str
    chunk_count: int
    created_at:  str

    @property
    def filename(self) -> str:
        return self.name

    @property
    def chunks(self) -> list:
        return []


# ══════════════════════════════════════════════════════════════════════════════
# Text extractors
# ══════════════════════════════════════════════════════════════════════════════

def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    pages  = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _extract_docx(content: bytes) -> str:
    from docx import Document as DocxDoc
    doc   = DocxDoc(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def _extract(filename: str, content: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return _extract_pdf(content)
    if ext == "docx":
        return _extract_docx(content)
    return content.decode("utf-8", errors="ignore")


# ══════════════════════════════════════════════════════════════════════════════
# Chunker
# ══════════════════════════════════════════════════════════════════════════════

CHUNK_SIZE    = 400
CHUNK_OVERLAP = 80


def _chunk_text(doc_id: str, doc_name: str, text: str) -> list[Chunk]:
    words  = text.split()
    step   = CHUNK_SIZE - CHUNK_OVERLAP
    chunks = []
    for i, start in enumerate(range(0, len(words), step)):
        chunk_text = " ".join(words[start : start + CHUNK_SIZE])
        if chunk_text.strip():
            chunks.append(
                Chunk(
                    id          = str(uuid.uuid4()),
                    doc_id      = doc_id,
                    doc_name    = doc_name,
                    chunk_index = i,
                    text        = chunk_text,
                )
            )
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# DocumentStore — SQLite-backed
# ══════════════════════════════════════════════════════════════════════════════

class DocumentStore:
    def __init__(self):
        self._db = DB_PATH
        self._init_db()

    def _conn(self):
        conn = sqlite3.connect(self._db)
        conn.row_factory = sqlite3.Row
        return conn

    def _init_db(self):
        with self._conn() as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id          TEXT PRIMARY KEY,
                    name        TEXT NOT NULL,
                    preview     TEXT NOT NULL,
                    chunk_count INTEGER NOT NULL,
                    created_at  TEXT NOT NULL
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS chunks (
                    id          TEXT PRIMARY KEY,
                    doc_id      TEXT NOT NULL,
                    doc_name    TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    text        TEXT NOT NULL,
                    FOREIGN KEY (doc_id) REFERENCES documents(id)
                )
            """)
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_chunks_doc ON chunks(doc_id)"
            )
            conn.commit()

    # ── internal ──────────────────────────────────────────────────────────────

    def _create(self, name: str, text: str) -> Document:
        doc_id = str(uuid.uuid4())
        chunks = _chunk_text(doc_id, name, text)

        doc = Document(
            id          = doc_id,
            name        = name,
            preview     = text[:200],
            chunk_count = len(chunks),
            created_at  = datetime.utcnow().isoformat(),
        )

        with self._conn() as conn:
            conn.execute(
                "INSERT INTO documents (id, name, preview, chunk_count, created_at) VALUES (?,?,?,?,?)",
                (doc.id, doc.name, doc.preview, doc.chunk_count, doc.created_at),
            )
            conn.executemany(
                "INSERT INTO chunks (id, doc_id, doc_name, chunk_index, text) VALUES (?,?,?,?,?)",
                [(c.id, c.doc_id, c.doc_name, c.chunk_index, c.text) for c in chunks],
            )
            conn.commit()

        return doc

    # ── add — new style ───────────────────────────────────────────────────────

    def add_document(self, name: str, content: str) -> Document:
        return self._create(name, content)

    # ── add — old style ───────────────────────────────────────────────────────

    def add_file(self, filename: str, content: bytes) -> Document:
        text = _extract(filename, content)
        return self._create(filename, text)

    def add_text(self, title: str, text: str) -> Document:
        return self._create(title, text)

    # ── remove ────────────────────────────────────────────────────────────────

    def remove_document(self, doc_id: str) -> bool:
        with self._conn() as conn:
            row = conn.execute(
                "SELECT id FROM documents WHERE id=?", (doc_id,)
            ).fetchone()
            if not row:
                return False
            conn.execute("DELETE FROM chunks WHERE doc_id=?", (doc_id,))
            conn.execute("DELETE FROM documents WHERE id=?", (doc_id,))
            conn.commit()
        return True

    def remove(self, doc_id: str) -> None:
        self.remove_document(doc_id)

    # ── query ─────────────────────────────────────────────────────────────────

    def list_documents(self) -> list[Document]:
        with self._conn() as conn:
            rows = conn.execute(
                "SELECT id, name, preview, chunk_count, created_at FROM documents ORDER BY created_at"
            ).fetchall()
        return [
            Document(
                id          = r["id"],
                name        = r["name"],
                preview     = r["preview"],
                chunk_count = r["chunk_count"],
                created_at  = r["created_at"],
            )
            for r in rows
        ]

    def get_chunks(self) -> list[Chunk]:
        with self._conn() as conn:
            rows = conn.execute(
                "SELECT id, doc_id, doc_name, chunk_index, text FROM chunks ORDER BY doc_id, chunk_index"
            ).fetchall()
        return [
            Chunk(
                id          = r["id"],
                doc_id      = r["doc_id"],
                doc_name    = r["doc_name"],
                chunk_index = r["chunk_index"],
                text        = r["text"],
            )
            for r in rows
        ]

    def total_chunks(self) -> int:
        with self._conn() as conn:
            row = conn.execute("SELECT COUNT(*) as n FROM chunks").fetchone()
        return row["n"]

    def get_filename(self, doc_id: str) -> str:
        with self._conn() as conn:
            row = conn.execute(
                "SELECT name FROM documents WHERE id=?", (doc_id,)
            ).fetchone()
        return row["name"] if row else "unknown"

    def stats(self) -> dict:
        with self._conn() as conn:
            doc_count   = conn.execute("SELECT COUNT(*) as n FROM documents").fetchone()["n"]
            chunk_count = conn.execute("SELECT COUNT(*) as n FROM chunks").fetchone()["n"]
        return {"documents": doc_count, "chunks": chunk_count}